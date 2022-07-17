from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fillincells():
    for k in range(len(recfixed)):
        row = table.add_row().cells
        for cells in table.rows[-1].cells:
            cells.text = lines[k]
            run = cells.paragraphs[0].runs[0]  # formatting occurs at run level, I'll call it as real-time
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.all_caps = True
            run.font.color.rgb = RGBColor(136, 8, 8)

def cellsmidifier():
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range((len(recfixed))):
        table.cell(0, 0 + i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 0 + i).vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
        print("Debug number first "+str(i))
        for j in range((len(recfixed)) - recfixindex):
            table.cell(0 + j, 0 + i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.cell(0 + j, 0 + i).vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            print("Debug number second " + str(j))
            for k in range((len(recfixed)) - recfixindex):
                table.cell(0 + j, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.cell(0 + j, 0).vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
                print("Debug number third" + str(k))

def testobject():
    print(lines)
    print("number of lines: " + str(count))
    print(recfixed)
    for i in range(len(recfixed)):
        print(i)
        i += 1

lines = []
with open('wnrslines.txt','r') as f:
    lines = f.readlines()

count = 0
for line in lines:
    count += 1

document = Document()

recfixed = tuple(lines)

testobject()

recfixindex = int((count / 2)-1)

table = document.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

fillincells()
for row in table.rows:
    row.height = Cm(5.6)

cellsmidifier()

# table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER #horizontal allignment
# table.cell(1, 1).vertical_alignment = WD_TABLE_ALIGNMENT.CENTER #vertical allignment

document.save('WNRSprint.docx')